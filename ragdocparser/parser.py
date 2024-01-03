"""
Document parser for various file formats including OCR support.
"""

import os
import logging
from pathlib import Path
from typing import Dict, List, Optional, Union
from abc import ABC, abstractmethod
from datetime import datetime

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from .ocr import OCRProcessor, ImageDocumentParser
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logger = logging.getLogger(__name__)


class BaseParser(ABC):
    """Base class for document parsers."""
    
    @abstractmethod
    def parse(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """Parse a document and return extracted content."""
        pass


class PDFParser(BaseParser):
    """Parser for PDF documents with OCR fallback."""
    
    def __init__(self, use_ocr_fallback: bool = True):
        self.use_ocr_fallback = use_ocr_fallback and OCR_AVAILABLE
        if self.use_ocr_fallback:
            self.ocr_processor = OCRProcessor()
    
    def parse(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """Parse PDF document with OCR fallback for image-based PDFs."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is required for PDF parsing")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        text_content = []
        metadata = {}
        extraction_method = "text_extraction"
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata safely
                if pdf_reader.metadata:
                    metadata.update({
                        'title': str(pdf_reader.metadata.get('/Title', '')),
                        'author': str(pdf_reader.metadata.get('/Author', '')),
                        'subject': str(pdf_reader.metadata.get('/Subject', '')),
                        'creator': str(pdf_reader.metadata.get('/Creator', '')),
                        'producer': str(pdf_reader.metadata.get('/Producer', '')),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_reader.metadata.get('/ModDate', '')),
                    })
                
                metadata['pages'] = len(pdf_reader.pages)
                
                # Try text extraction first
                total_text_length = 0
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            cleaned_text = ' '.join(page_text.split())
                            text_content.append({
                                'page': page_num + 1,
                                'content': cleaned_text
                            })
                            total_text_length += len(cleaned_text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
                
                # If very little text was extracted, try OCR
                if (total_text_length < 100 and self.use_ocr_fallback):
                    logger.info(f"Low text extraction yield ({total_text_length} chars), trying OCR...")
                    try:
                        ocr_results = self.ocr_processor.extract_text_from_pdf_images(file_path)
                        
                        # Replace with OCR results if they contain more text
                        ocr_text_length = sum(len(result.get('text', '')) for result in ocr_results)
                        
                        if ocr_text_length > total_text_length:
                            text_content = []
                            for result in ocr_results:
                                if result.get('text', '').strip():
                                    text_content.append({
                                        'page': result.get('page', 1),
                                        'content': result['text'].strip()
                                    })
                            
                            extraction_method = "OCR"
                            metadata['ocr_confidence'] = sum(r.get('confidence', 0) for r in ocr_results) / len(ocr_results)
                            logger.info(f"OCR extracted {ocr_text_length} characters vs {total_text_length} from text extraction")
                    
                    except Exception as e:
                        logger.warning(f"OCR fallback failed: {e}")
                        
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
        
        if not text_content:
            logger.warning(f"No text content extracted from PDF: {file_path}")
        
        metadata['extraction_method'] = extraction_method
        
        return {
            'content': text_content,
            'metadata': metadata,
            'file_path': str(file_path),
            'file_type': 'pdf'
        }


class TXTParser(BaseParser):
    """Parser for plain text documents."""
    
    def parse(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """Parse text document with encoding detection."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        content = None
        encoding_used = None
        
        # Try multiple encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    encoding_used = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError(f"Unable to decode file {file_path} with any supported encoding")
        
        # Get file statistics
        stat = file_path.stat()
        metadata = {
            'file_size': stat.st_size,
            'creation_date': stat.st_ctime,
            'modification_date': stat.st_mtime,
            'encoding': encoding_used,
            'lines': len(content.splitlines()),
            'characters': len(content),
            'words': len(content.split()),
            'extraction_method': 'direct_read'
        }
        
        return {
            'content': [{'page': 1, 'content': content.strip()}],
            'metadata': metadata,
            'file_path': str(file_path),
            'file_type': 'txt'
        }


class DocxParser(BaseParser):
    """Parser for DOCX documents."""
    
    def parse(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """Parse DOCX document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing")
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text content
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            content = '
'.join(paragraphs)
            
            # Extract metadata
            props = doc.core_properties
            metadata = {
                'title': props.title or '',
                'author': props.author or '',
                'subject': props.subject or '',
                'keywords': props.keywords or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else '',
                'last_modified_by': props.last_modified_by or '',
                'paragraphs': len(doc.paragraphs),
                'words': len(content.split()) if content else 0,
                'extraction_method': 'docx_parsing'
            }
            
            return {
                'content': [{'page': 1, 'content': content}],
                'metadata': metadata,
                'file_path': str(file_path),
                'file_type': 'docx'
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise


class DocumentParser:
    """Main document parser that handles multiple file formats including OCR."""
    
    def __init__(self, use_ocr: bool = True):
        self.use_ocr = use_ocr and OCR_AVAILABLE
        
        self.parsers = {
            '.pdf': PDFParser(use_ocr_fallback=self.use_ocr),
            '.txt': TXTParser(),
        }
        
        if DOCX_AVAILABLE:
            self.parsers['.docx'] = DocxParser()
        
        # Add image parsers if OCR is available
        if self.use_ocr:
            image_parser = ImageDocumentParser()
            for fmt in image_parser.supported_formats:
                self.parsers[fmt] = image_parser
    
    def parse(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """Parse a document based on its file extension."""
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension not in self.parsers:
            raise ValueError(f"Unsupported file format: {extension}. Supported formats: {self.supported_formats()}")
        
        parser = self.parsers[extension]
        result = parser.parse(file_path)
        
        # Add common metadata
        stat = file_path.stat()
        result['metadata'].update({
            'file_name': file_path.name,
            'file_size_bytes': stat.st_size,
            'parsed_at': str(datetime.now()),
        })
        
        return result
    
    def supported_formats(self) -> List[str]:
        """Return list of supported file formats."""
        return list(self.parsers.keys())
    
    def parse_directory(self, directory_path: Union[str, Path], recursive: bool = True) -> List[Dict[str, any]]:
        """Parse all supported documents in a directory."""
        directory_path = Path(directory_path)
        if not directory_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        results = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory_path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.parsers:
                try:
                    result = self.parse(file_path)
                    results.append(result)
                    logger.info(f"Successfully parsed: {file_path}")
                except Exception as e:
                    logger.error(f"Failed to parse {file_path}: {e}")
                    continue
        
        return results
